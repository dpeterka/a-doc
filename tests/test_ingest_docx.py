"""Tests for adoc.ingest.docx.extract_docx_text: deterministic, no LLM."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from adoc.ingest.docx import DocxExtractionError, extract_docx_text


def test_extracts_paragraphs_in_order_and_skips_blanks(tmp_path: Path) -> None:
    path = tmp_path / "narrative.docx"
    document = Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("")  # blank - must be skipped
    document.add_paragraph("   ")  # whitespace-only - must be skipped
    document.add_paragraph("Second paragraph.")
    document.save(str(path))

    text = extract_docx_text(path)

    assert text == "First paragraph.\n\nSecond paragraph."


def test_extraction_is_deterministic(tmp_path: Path) -> None:
    path = tmp_path / "narrative.docx"
    document = Document()
    document.add_paragraph("Same every time.")
    document.save(str(path))

    assert extract_docx_text(path) == extract_docx_text(path)


def test_extracts_a_table_as_pipe_rows_in_document_order(tmp_path: Path) -> None:
    path = tmp_path / "labs.docx"
    document = Document()
    document.add_paragraph("Lab panel:")
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Test"
    table.cell(0, 1).text = "Result"
    table.cell(0, 2).text = "Range"
    table.cell(1, 0).text = "Potassium"
    table.cell(1, 1).text = "4.1"
    table.cell(1, 2).text = "3.5-5.1"
    document.add_paragraph("End of panel.")
    document.save(str(path))

    text = extract_docx_text(path)

    assert "| Test | Result | Range |" in text
    assert "| Potassium | 4.1 | 3.5-5.1 |" in text
    # document order preserved: paragraph, then table, then paragraph
    assert text.index("Lab panel:") < text.index("| Test") < text.index("End of panel.")


def test_skips_fully_blank_table_rows(tmp_path: Path) -> None:
    path = tmp_path / "labs.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Test"
    table.cell(0, 1).text = "Result"
    # row 1 left entirely blank
    document.save(str(path))

    text = extract_docx_text(path)

    assert text == "| Test | Result |"


def test_raises_docx_extraction_error_on_unreadable_file(tmp_path: Path) -> None:
    path = tmp_path / "bogus.docx"
    path.write_bytes(b"not a docx package at all")

    with pytest.raises(DocxExtractionError):
        extract_docx_text(path)
