"""Tests for adoc.ingest.filetypes.detect_doc_kind."""

from __future__ import annotations

from pathlib import Path

from conftest import TINY_PDF_BYTES
from docx import Document

from adoc.ingest.filetypes import detect_doc_kind


def test_detects_pdf_by_magic(tmp_path: Path) -> None:
    path = tmp_path / "doc.pdf"
    path.write_bytes(TINY_PDF_BYTES)
    assert detect_doc_kind(path) == "pdf"


def test_detects_a_real_docx_package(tmp_path: Path) -> None:
    path = tmp_path / "notes.docx"
    document = Document()
    document.add_paragraph("hello")
    document.save(str(path))

    assert detect_doc_kind(path) == "docx"


def test_rejects_a_fake_pk_header_without_a_real_zip(tmp_path: Path) -> None:
    """A file merely renamed to `.docx` with a spoofed zip magic, but not an
    actual zip archive, must never be treated as a real docx package."""
    path = tmp_path / "notes.docx"
    path.write_bytes(b"PK\x03\x04 not a real zip at all")

    assert detect_doc_kind(path) is None


def test_rejects_a_real_zip_without_content_types(tmp_path: Path) -> None:
    """A real zip archive that isn't an OOXML package (no
    `[Content_Types].xml`) must not be misdetected as docx even with the
    `.docx` suffix."""
    import zipfile

    path = tmp_path / "notes.docx"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("readme.txt", "just a plain zip")

    assert detect_doc_kind(path) is None


def test_rejects_a_real_docx_package_with_the_wrong_suffix(tmp_path: Path) -> None:
    """The suffix check is deliberate - a renamed-away `.docx` is out of
    scope, matching the pdf path's own filename-agnostic-but-magic-gated
    behavior (the magic bytes alone are the zip format's, shared by many
    non-docx formats)."""
    path = tmp_path / "notes.zip"
    document = Document()
    document.add_paragraph("hello")
    document.save(str(path))

    assert detect_doc_kind(path) is None


def test_rejects_unrelated_file_types(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("just some text", encoding="utf-8")

    assert detect_doc_kind(path) is None


def test_rejects_a_missing_file_without_raising(tmp_path: Path) -> None:
    assert detect_doc_kind(tmp_path / "missing.pdf") is None
