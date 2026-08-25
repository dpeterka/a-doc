"""Tests for adoc.ingest.doctext: the document-TEXT layer
(docs/adr/0015-document-text-corpus.md) — extraction per kind, the genomics
exclusion, storage/backfill idempotency, and rebuild-from-committed-files
coherence.

PDF text extraction is always exercised via an injected `pdf_extractor`
fake — never the real `pdftotext` binary — so these tests never depend on
poppler being installed (mirrors `ingest.archive`'s injected
`PageRenderer`).
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from adoc.casefile.repo import DataRepo
from adoc.ingest.doctext import (
    backfill_document_text,
    document_text_path,
    extract_text_for_kind,
    rebuild_document_text_from_files,
    store_document_text,
)
from adoc.ingest.genomics import GENOMIC_DOC_TYPE, archive_genomic_file
from adoc.labs.db import LabsDb
from adoc.labs.models import DocumentStatus, LabDocument

SHA = "d" * 64


@pytest.fixture
def repo(tmp_path: Path) -> DataRepo:
    return DataRepo.init_at(tmp_path / "data")


@pytest.fixture
def db(tmp_path: Path) -> LabsDb:
    return LabsDb(tmp_path / "labs.sqlite")


def _fake_pdf_extractor(text: str | None):  # type: ignore[no-untyped-def]
    def extractor(path: Path) -> str | None:
        return text

    return extractor


def _upsert_doc(db: LabsDb, sha: str, *, doc_type: str = "lab_report") -> None:
    db.upsert_document(
        LabDocument(
            sha256=sha,
            filename="report.pdf",
            doc_type=doc_type,
            doc_date=date(2026, 5, 2),
            page_count=1,
            status=DocumentStatus.COMPLETE,
        )
    )


# --------------------------------------------------------------------------
# extract_text_for_kind — per type
# --------------------------------------------------------------------------


def test_extract_text_for_kind_pdf_uses_injected_extractor(tmp_path: Path) -> None:
    pdf_path = tmp_path / "doc.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n")
    text = extract_text_for_kind(
        "pdf", pdf_path, pdf_extractor=_fake_pdf_extractor("Impression: normal.\fPage two text.")
    )
    assert text == "Impression: normal.\fPage two text."


def test_extract_text_for_kind_pdf_extractor_failure_returns_none(tmp_path: Path) -> None:
    pdf_path = tmp_path / "doc.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n")
    text = extract_text_for_kind("pdf", pdf_path, pdf_extractor=_fake_pdf_extractor(None))
    assert text is None


def test_extract_text_for_kind_docx(tmp_path: Path) -> None:
    docx_path = tmp_path / "history.docx"
    document = Document()
    document.add_paragraph("Patient-authored clinical history.")
    document.save(str(docx_path))
    text = extract_text_for_kind("docx", docx_path)
    assert text is not None
    assert "Patient-authored clinical history." in text


def test_extract_text_for_kind_text_reads_verbatim(tmp_path: Path) -> None:
    txt_path = tmp_path / "notes.txt"
    txt_path.write_text("plain notes here", encoding="utf-8")
    assert extract_text_for_kind("text", txt_path) == "plain notes here"


# --------------------------------------------------------------------------
# store_document_text: pagination + persistence
# --------------------------------------------------------------------------


def test_store_document_text_writes_committed_file_and_db_rows(repo: DataRepo, db: LabsDb) -> None:
    _upsert_doc(db, SHA)
    path = store_document_text(repo, db, SHA, "page one\fpage two")

    assert path == document_text_path(repo.root, SHA)
    assert path.read_text(encoding="utf-8") == "page one\fpage two"
    assert db.document_text_shas() == {SHA}
    assert db.get_document_text(SHA) == "page one\fpage two"

    # "one" (not "page") is the query, so this is unambiguous under
    # `search_document_text`'s OR-of-tokens ranking (both pages contain
    # "page"; only page 1 contains "one").
    hits = db.search_document_text("one")
    assert len(hits) == 1
    assert hits[0].page == 1
    assert hits[0].source_ref == "doc:report.pdf#p1"


def test_store_document_text_no_form_feed_is_page_none(repo: DataRepo, db: LabsDb) -> None:
    _upsert_doc(db, SHA)
    store_document_text(repo, db, SHA, "a single unpaginated blob of text")
    hits = db.search_document_text("unpaginated")
    assert hits[0].page is None
    assert hits[0].source_ref == "doc:report.pdf"


def test_store_document_text_is_idempotent(repo: DataRepo, db: LabsDb) -> None:
    _upsert_doc(db, SHA)
    store_document_text(repo, db, SHA, "first version")
    store_document_text(repo, db, SHA, "second version")
    assert db.get_document_text(SHA) == "second version"
    path = document_text_path(repo.root, SHA)
    assert path.read_text(encoding="utf-8") == "second version"


# --------------------------------------------------------------------------
# rebuild_document_text_from_files: derived-from-committed-files coherence
# --------------------------------------------------------------------------


def test_rebuild_from_files_reproduces_the_same_pages(repo: DataRepo, db: LabsDb) -> None:
    _upsert_doc(db, SHA)
    store_document_text(repo, db, SHA, "page one\fpage two")

    # Simulate a fresh checkout: a brand-new (empty) sqlite db, same repo files.
    fresh_db = LabsDb(":memory:")
    _upsert_doc(fresh_db, SHA)
    count = rebuild_document_text_from_files(fresh_db, repo.root)

    assert count == 1
    assert fresh_db.get_document_text(SHA) == "page one\fpage two"
    hits = fresh_db.search_document_text("page two")
    assert hits[0].page == 2


def test_rebuild_from_files_skips_unknown_shas(repo: DataRepo, db: LabsDb) -> None:
    _upsert_doc(db, SHA)
    store_document_text(repo, db, SHA, "some text")

    fresh_db = LabsDb(":memory:")  # documents table empty - SHA is unknown
    count = rebuild_document_text_from_files(fresh_db, repo.root)
    assert count == 0
    assert fresh_db.document_text_shas() == set()


def test_rebuild_from_files_no_dir_returns_zero(repo: DataRepo, db: LabsDb) -> None:
    assert rebuild_document_text_from_files(db, repo.root) == 0


# --------------------------------------------------------------------------
# backfill_document_text: coverage sweep + idempotency
# --------------------------------------------------------------------------


def test_backfill_extracts_missing_text_for_docx_and_text_documents(
    repo: DataRepo, db: LabsDb
) -> None:
    docx_sha = "1" * 64
    txt_sha = "2" * 64

    docx_original = repo.root / "sources" / f"{docx_sha}__history.docx"
    document = Document()
    document.add_paragraph("Onset of joint pain in March.")
    document.save(str(docx_original))
    _upsert_doc(db, docx_sha, doc_type="clinical_note")
    db._conn.execute(
        "UPDATE documents SET filename = ? WHERE sha256 = ?", ("history.docx", docx_sha)
    )

    txt_original = repo.root / "sources" / f"{txt_sha}__note.txt"
    txt_original.write_text("A plain text note.", encoding="utf-8")
    _upsert_doc(db, txt_sha, doc_type="other")
    db._conn.execute("UPDATE documents SET filename = ? WHERE sha256 = ?", ("note.txt", txt_sha))
    db._conn.commit()

    report = backfill_document_text(repo, db)

    assert report.total_non_genomic == 2
    assert report.already_covered == 0
    assert report.extracted == 2
    assert report.skipped_no_source == 0
    assert report.skipped_genomic == 0

    assert "Onset of joint pain in March." in (db.get_document_text(docx_sha) or "")
    assert db.get_document_text(txt_sha) == "A plain text note."


def test_backfill_is_idempotent_second_run_extracts_nothing(repo: DataRepo, db: LabsDb) -> None:
    txt_sha = "3" * 64
    txt_original = repo.root / "sources" / f"{txt_sha}__note.txt"
    txt_original.write_text("content", encoding="utf-8")
    _upsert_doc(db, txt_sha, doc_type="other")
    db._conn.execute("UPDATE documents SET filename = ? WHERE sha256 = ?", ("note.txt", txt_sha))
    db._conn.commit()

    first = backfill_document_text(repo, db)
    assert first.extracted == 1

    second = backfill_document_text(repo, db)
    assert second.extracted == 0
    assert second.already_covered == 1


def test_backfill_skips_document_with_no_archived_source(repo: DataRepo, db: LabsDb) -> None:
    missing_sha = "4" * 64
    _upsert_doc(db, missing_sha, doc_type="other")  # no file under sources/ for this sha
    report = backfill_document_text(repo, db)
    assert report.total_non_genomic == 1
    assert report.extracted == 0
    assert report.skipped_no_source == 1


# --------------------------------------------------------------------------
# Genomics exclusion (CRITICAL DESIGN RULE, ADR 0010 — never weakened here)
# --------------------------------------------------------------------------


def test_genomic_documents_are_never_extracted(repo: DataRepo, db: LabsDb) -> None:
    """A genomic file, archived the real way (`ingest.genomics.
    archive_genomic_file`, never `archive_document`), must never gain a
    `document_text` row via `adoc backfill-doc-text` — proven with a
    `pdf_extractor` stub that raises if ever invoked, mirroring
    `test_ingest_pipeline.py`'s "exploding vision client" pattern for the
    same guarantee at the live-ingest layer.
    """

    def exploding_pdf_extractor(path: Path) -> str | None:
        raise AssertionError("pdf_extractor must never be called for a genomic file")

    genomic_path = repo.root.parent / "23andme_raw.txt"
    genomic_path.write_text("# comment\nrsid\tchromosome\tposition\tgenotype\n", encoding="utf-8")
    archived = archive_genomic_file(repo.root, genomic_path, db=db)
    db.upsert_document(
        LabDocument(
            sha256=archived.sha256,
            filename="23andme_raw.txt",
            doc_type=GENOMIC_DOC_TYPE,
            doc_date=None,
            page_count=1,
            status=DocumentStatus.COMPLETE,
        )
    )

    # Also seed one genuinely-extractable document so the sweep does real work.
    txt_sha = "5" * 64
    txt_original = repo.root / "sources" / f"{txt_sha}__ordinary.txt"
    txt_original.write_text("ordinary document text", encoding="utf-8")
    _upsert_doc(db, txt_sha, doc_type="other")
    db._conn.execute(
        "UPDATE documents SET filename = ? WHERE sha256 = ?", ("ordinary.txt", txt_sha)
    )
    db._conn.commit()

    report = backfill_document_text(repo, db, pdf_extractor=exploding_pdf_extractor)

    assert report.skipped_genomic == 1
    assert report.total_non_genomic == 1  # the genomic doc is not counted here
    assert report.extracted == 1  # only the ordinary text document
    assert archived.sha256 not in db.document_text_shas()
    assert db.get_document_text(archived.sha256) is None


def test_extract_text_for_kind_type_excludes_genomic() -> None:
    """Type-level proof (module docstring): `DocKind` — the only type
    `extract_text_for_kind`'s `kind` parameter accepts — has no `"genomic"`
    member at all, so no genomic file's kind can ever reach this function
    in the first place."""
    from adoc.ingest.archive import DocKind

    assert "genomic" not in DocKind.__args__  # type: ignore[attr-defined]
