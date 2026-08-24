"""Tests for adoc.ingest.pipeline: end-to-end wiring over fixture extractions.

Fake `VisionClient`s return fixture `DocumentExtraction`/`ClassifyResult`
payloads (`tests/fixtures/extractions/*.json`) instead of calling any real
provider - no network in tests. A fake page renderer stands in for
`pdftoppm` (no poppler in CI).
"""

from __future__ import annotations

import json
from collections.abc import Sequence
from datetime import UTC, date, datetime
from pathlib import Path
from typing import Any

from conftest import TINY_PDF_BYTES, fake_page_renderer
from docx import Document
from git import Repo
from pydantic import BaseModel

from adoc.casefile.encounters import read_encounter
from adoc.casefile.repo import DataRepo
from adoc.config import ModelBinding
from adoc.ingest.pipeline import ingest_directory, ingest_file, ingest_inbox
from adoc.ingest.schema import ClassifyResult, DocumentExtraction
from adoc.ingest.vision import Part, VisionClient
from adoc.labs.db import LabsDb
from adoc.labs.models import DocumentStatus
from adoc.reason.client import (
    AnthropicProvider,
    LlmClient,
    OpenAIProvider,
    TransportRequest,
    TransportResponse,
)

FIXTURES_DIR = Path(__file__).parent / "fixtures" / "extractions"


def _load_fixture(name: str) -> tuple[DocumentExtraction, DocumentExtraction]:
    payload = json.loads((FIXTURES_DIR / name).read_text(encoding="utf-8"))
    return (
        DocumentExtraction.model_validate(payload["pass_a"]),
        DocumentExtraction.model_validate(payload["pass_b"]),
    )


class FakeVisionClient:
    """Stands in for `VisionClient`: routes by role to fixture payloads."""

    def __init__(self, fixture_name: str) -> None:
        self.pass_a, self.pass_b = _load_fixture(fixture_name)
        self.calls: list[str] = []

    def extract(
        self,
        role: str,
        *,
        system: str,
        parts: Sequence[Part],
        schema: type[BaseModel],
        binding_index: int = 0,
        max_tokens: int = 4096,
    ) -> Any:
        self.calls.append(role)
        if role == "classifier":
            doc_date = self.pass_a.collection_date or self.pass_a.report_date
            return ClassifyResult(doc_type=self.pass_a.doc_type, doc_date=doc_date)
        if role == "extractor_pass_a":
            return self.pass_a
        if role == "extractor_pass_b":
            return self.pass_b
        raise AssertionError(f"unexpected role: {role}")


def _fixed_clock() -> datetime:
    return datetime(2026, 5, 3, 12, 0, 0, tzinfo=UTC)


def _ingest(path: Path, *, repo: DataRepo, db: LabsDb, vision: FakeVisionClient):  # type: ignore[no-untyped-def]
    return ingest_file(
        path,
        repo=repo,
        db=db,
        vision=vision,  # type: ignore[arg-type]
        clock=_fixed_clock,
        renderer=fake_page_renderer(1),
    )


def test_lab_report_full_agreement_is_all_auto_and_exports_jsonl(tmp_path: Path) -> None:
    repo = DataRepo.init_at(tmp_path / "data-repo")
    db = LabsDb(tmp_path / "labs.sqlite")
    doc_path = tmp_path / "quest-2026-05-02.pdf"
    doc_path.write_bytes(TINY_PDF_BYTES)
    vision = FakeVisionClient("clean_agreement.json")

    report = ingest_file(
        doc_path,
        repo=repo,
        db=db,
        vision=vision,  # type: ignore[arg-type]
        clock=_fixed_clock,
        renderer=fake_page_renderer(1),
    )

    assert len(report.files) == 1
    outcome = report.files[0]
    assert outcome.outcome == "ingested"
    assert outcome.doc_type == "lab_report"
    assert outcome.rows_auto == 2
    assert outcome.rows_pending == 0
    assert outcome.issues == []
    assert outcome.commit_sha is not None

    export_path = repo.root / "labs-export.jsonl"
    assert export_path.exists()
    lines = export_path.read_text(encoding="utf-8").strip().splitlines()
    assert any(json.loads(line)["table"] == "lab" for line in lines)

    documents = db.list_documents()
    assert len(documents) == 1
    assert documents[0].status == DocumentStatus.COMPLETE
    assert documents[0].doc_date == date(2026, 5, 2)


def test_commit_message_has_expected_shape(tmp_path: Path) -> None:
    repo = DataRepo.init_at(tmp_path / "data-repo")
    db = LabsDb(tmp_path / "labs.sqlite")
    doc_path = tmp_path / "quest.pdf"
    doc_path.write_bytes(TINY_PDF_BYTES)
    vision = FakeVisionClient("clean_agreement.json")

    _ingest(doc_path, repo=repo, db=db, vision=vision)

    git_repo = Repo(repo.root)
    message = git_repo.head.commit.message
    assert message.startswith("ingest: Quest Diagnostics 2026-05-02")
    assert "(2 rows, 0 queued)" in message


def test_ingesting_the_same_document_twice_dedupes_and_makes_no_second_commit(
    tmp_path: Path,
) -> None:
    repo = DataRepo.init_at(tmp_path / "data-repo")
    db = LabsDb(tmp_path / "labs.sqlite")
    doc_path = tmp_path / "quest.pdf"
    doc_path.write_bytes(TINY_PDF_BYTES)
    vision = FakeVisionClient("clean_agreement.json")

    first = _ingest(doc_path, repo=repo, db=db, vision=vision)
    commit_count_after_first = len(list(Repo(repo.root).iter_commits()))

    second = _ingest(doc_path, repo=repo, db=db, vision=vision)
    commit_count_after_second = len(list(Repo(repo.root).iter_commits()))

    assert first.files[0].outcome == "ingested"
    assert second.files[0].outcome == "duplicate"
    assert second.files[0].commit_sha is None
    assert commit_count_after_second == commit_count_after_first
    # the classifier/extractor roles must not even be called on the duplicate
    assert vision.calls.count("classifier") == 1


def test_value_disagreement_queues_the_row_pending(tmp_path: Path) -> None:
    repo = DataRepo.init_at(tmp_path / "data-repo")
    db = LabsDb(tmp_path / "labs.sqlite")
    doc_path = tmp_path / "quest.pdf"
    doc_path.write_bytes(TINY_PDF_BYTES)
    vision = FakeVisionClient("value_disagreement.json")

    report = _ingest(doc_path, repo=repo, db=db, vision=vision)

    outcome = report.files[0]
    assert outcome.rows_auto == 0
    assert outcome.rows_pending == 1
    assert any("value_mismatch" in issue for issue in outcome.issues)

    pending_rows = db.pending()
    assert len(pending_rows) == 1
    assert pending_rows[0].name == "potassium"


def test_non_lab_document_creates_an_encounter_stub_not_a_labs_row(tmp_path: Path) -> None:
    repo = DataRepo.init_at(tmp_path / "data-repo")
    db = LabsDb(tmp_path / "labs.sqlite")
    doc_path = tmp_path / "rheum-visit.pdf"
    doc_path.write_bytes(TINY_PDF_BYTES)
    vision = FakeVisionClient("non_lab_clinical_note.json")

    report = _ingest(doc_path, repo=repo, db=db, vision=vision)

    outcome = report.files[0]
    assert outcome.outcome == "ingested"
    assert outcome.doc_type == "clinical_note"
    assert outcome.rows_auto == 0
    assert outcome.rows_pending == 0
    # extractor roles are never called for a non-lab document
    assert "extractor_pass_a" not in vision.calls
    assert "extractor_pass_b" not in vision.calls
    assert db.pending() == []

    encounters_dir = repo.root / "case" / "encounters"
    encounter_files = list(encounters_dir.glob("2026-06-10--*.md"))
    assert len(encounter_files) == 1
    encounter = read_encounter(encounter_files[0])
    assert encounter.frontmatter.type == "specialist-visit"
    assert encounter.summary == "(pending review)"
    assert encounter.frontmatter.sources


def test_ingest_inbox_scans_and_ingests_every_file(tmp_path: Path) -> None:
    repo = DataRepo.init_at(tmp_path / "data-repo")
    db = LabsDb(tmp_path / "labs.sqlite")
    (repo.root / "inbox" / "a.pdf").write_bytes(TINY_PDF_BYTES)
    (repo.root / "inbox" / "b.pdf").write_bytes(TINY_PDF_BYTES + b"\x00")
    vision = FakeVisionClient("clean_agreement.json")

    report = ingest_inbox(
        repo=repo,
        db=db,
        vision=vision,
        clock=_fixed_clock,
        renderer=fake_page_renderer(1),  # type: ignore[arg-type]
    )

    assert len(report.files) == 2
    assert all(f.outcome == "ingested" for f in report.files)
    assert report.total_auto == 4
    assert report.total_pending == 0


def test_scan_files_is_recursive_and_skips_sync_artifacts(tmp_path: Path) -> None:
    """Dropbox/rclone preserve subfolders in the inbox; nested files must be found."""
    from adoc.ingest.pipeline import _scan_files

    (tmp_path / "Labs" / "LabCorp").mkdir(parents=True)
    (tmp_path / "Labs" / "LabCorp" / "b.pdf").write_bytes(b"x")
    (tmp_path / "a.pdf").write_bytes(b"x")
    (tmp_path / "desktop.ini").write_bytes(b"x")
    (tmp_path / ".hidden").write_bytes(b"x")

    names = [p.name for p in _scan_files(tmp_path)]

    # Deterministic full-path order: "Labs/LabCorp/b.pdf" sorts before "a.pdf".
    assert names == ["b.pdf", "a.pdf"]


# --------------------------------------------------------------------------
# Post-ingest inbox hygiene (pipeline module docstring): `ingest_inbox`
# (and any `ingest_file` call that opts in via `inbox_root=`) deletes an
# ingested/duplicate inbox file and moves an error'd one to `work/failed/`
# with a `failures.jsonl` record. `ingest_directory` (the `adoc backfill
# <external dir>` engine) never applies this - the invariant tests below
# pin that a backfilled directory is never touched, success or failure.
# --------------------------------------------------------------------------


def test_ingest_inbox_deletes_the_file_once_ingested(tmp_path: Path) -> None:
    repo = DataRepo.init_at(tmp_path / "data-repo")
    db = LabsDb(tmp_path / "labs.sqlite")
    inbox_file = repo.root / "inbox" / "quest.pdf"
    inbox_file.write_bytes(TINY_PDF_BYTES)
    vision = FakeVisionClient("clean_agreement.json")

    report = ingest_inbox(
        repo=repo, db=db, vision=vision, clock=_fixed_clock, renderer=fake_page_renderer(1)
    )

    assert report.files[0].outcome == "ingested"
    assert not inbox_file.exists()
    # The immutable archive copy is authoritative and untouched.
    assert any((repo.root / "sources").glob(f"{report.files[0].sha256}__quest.pdf"))


def test_ingest_inbox_deletes_the_file_once_duplicate(tmp_path: Path) -> None:
    repo = DataRepo.init_at(tmp_path / "data-repo")
    db = LabsDb(tmp_path / "labs.sqlite")
    inbox_dir = repo.root / "inbox"
    (inbox_dir / "first.pdf").write_bytes(TINY_PDF_BYTES)
    vision = FakeVisionClient("clean_agreement.json")
    ingest_inbox(
        repo=repo, db=db, vision=vision, clock=_fixed_clock, renderer=fake_page_renderer(1)
    )

    # Identical bytes, dropped into the inbox again under a new name.
    (inbox_dir / "second.pdf").write_bytes(TINY_PDF_BYTES)
    report = ingest_inbox(
        repo=repo, db=db, vision=vision, clock=_fixed_clock, renderer=fake_page_renderer(1)
    )

    assert report.files[0].outcome == "duplicate"
    assert not (inbox_dir / "second.pdf").exists()


def test_ingest_inbox_moves_a_failed_file_to_work_failed_with_a_record(tmp_path: Path) -> None:
    repo = DataRepo.init_at(tmp_path / "data-repo")
    db = LabsDb(tmp_path / "labs.sqlite")
    inbox_dir = repo.root / "inbox"
    bad = inbox_dir / "corrupt.pdf"
    bad.write_bytes(b"not really a pdf")
    vision = FakeVisionClient("clean_agreement.json")

    report = ingest_inbox(
        repo=repo, db=db, vision=vision, clock=_fixed_clock, renderer=fake_page_renderer(1)
    )

    assert report.files[0].outcome == "error"
    assert not bad.exists()

    failed_path = repo.root / "work" / "failed" / "corrupt.pdf"
    assert failed_path.exists()
    assert failed_path.read_bytes() == b"not really a pdf"

    failures_log = repo.root / "work" / "failed" / "failures.jsonl"
    lines = failures_log.read_text(encoding="utf-8").strip().splitlines()
    assert len(lines) == 1
    record = json.loads(lines[0])
    assert record["filename"] == "corrupt.pdf"
    assert record["original_inbox_path"] == "corrupt.pdf"
    assert record["reason"]
    # Pydantic serializes a UTC datetime with a "Z" suffix rather than
    # "+00:00" - compare the parsed values, not the raw strings.
    assert datetime.fromisoformat(record["failed_at"]) == _fixed_clock()


def test_ingest_inbox_flattens_nested_failed_paths(tmp_path: Path) -> None:
    repo = DataRepo.init_at(tmp_path / "data-repo")
    db = LabsDb(tmp_path / "labs.sqlite")
    nested = repo.root / "inbox" / "Labs" / "LabCorp"
    nested.mkdir(parents=True)
    bad = nested / "b.pdf"
    bad.write_bytes(b"still not a pdf")
    vision = FakeVisionClient("clean_agreement.json")

    ingest_inbox(
        repo=repo, db=db, vision=vision, clock=_fixed_clock, renderer=fake_page_renderer(1)
    )

    flattened = repo.root / "work" / "failed" / "Labs__LabCorp__b.pdf"
    assert flattened.exists()
    record = json.loads(
        (repo.root / "work" / "failed" / "failures.jsonl").read_text(encoding="utf-8").strip()
    )
    assert record["original_inbox_path"] == "Labs/LabCorp/b.pdf"


def test_backfill_directory_never_deletes_or_moves_user_files_even_on_error(
    tmp_path: Path,
) -> None:
    """`adoc backfill <external dir>` must never touch the patient's own
    files - success or failure."""
    repo = DataRepo.init_at(tmp_path / "data-repo")
    db = LabsDb(tmp_path / "labs.sqlite")
    external_dir = tmp_path / "external-photos"
    external_dir.mkdir()
    bad_file = external_dir / "notes.txt"
    original_bytes = b"just some text, not a real document"
    bad_file.write_bytes(original_bytes)
    vision = FakeVisionClient("clean_agreement.json")

    report = ingest_directory(
        external_dir,
        repo=repo,
        db=db,
        vision=vision,
        clock=_fixed_clock,
        renderer=fake_page_renderer(1),
    )

    assert report.files[0].outcome == "error"
    assert bad_file.exists()
    assert bad_file.read_bytes() == original_bytes
    assert not (repo.root / "work" / "failed").exists()


def test_backfill_directory_never_deletes_a_successfully_ingested_file(tmp_path: Path) -> None:
    """Same invariant, the success path: backfilling still never deletes
    the original file (unlike `ingest_inbox`, which does)."""
    repo = DataRepo.init_at(tmp_path / "data-repo")
    db = LabsDb(tmp_path / "labs.sqlite")
    external_dir = tmp_path / "external-photos"
    external_dir.mkdir()
    good_file = external_dir / "quest.pdf"
    good_file.write_bytes(TINY_PDF_BYTES)
    vision = FakeVisionClient("clean_agreement.json")

    report = ingest_directory(
        external_dir,
        repo=repo,
        db=db,
        vision=vision,
        clock=_fixed_clock,
        renderer=fake_page_renderer(1),
    )

    assert report.files[0].outcome == "ingested"
    assert good_file.exists()
    assert good_file.read_bytes() == TINY_PDF_BYTES


# --------------------------------------------------------------------------
# docx routing (PLAN.md docx ingestion: docx = TEXT documents). These build
# a real `.docx` with `python-docx` and a real `VisionClient` wrapping a
# fake-transport `LlmClient` (mirroring `tests/test_stages.py`'s pattern) so
# `_ingest_docx`'s actual seam - `vision.client` - is exercised for real,
# never a duck-typed stand-in.
# --------------------------------------------------------------------------


def _make_narrative_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Patient-authored clinical history.")
    document.add_paragraph("Onset of joint pain in March 2026, worsening through the summer.")
    document.save(str(path))


def _make_lab_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Home lab panel, ordered 2026-08-01.")
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Test"
    table.cell(0, 1).text = "Result"
    table.cell(0, 2).text = "Range"
    table.cell(1, 0).text = "Potassium"
    table.cell(1, 1).text = "4.1"
    table.cell(1, 2).text = "3.5-5.1"
    document.save(str(path))


def _build_docx_llm_client(
    *,
    classify_payload: dict[str, Any],
    pass_a_payload: dict[str, Any] | None,
    pass_b_payload: dict[str, Any] | None,
    calls: list[str],
) -> LlmClient:
    """`classifier`/`extractor_pass_a` -> anthropic, `extractor_pass_b` ->
    openai (mirrors `models.yaml`'s real bindings); the anthropic transport
    dispatches on the requested schema since both roles share that
    provider."""

    def anthropic_transport(request: TransportRequest) -> TransportResponse:
        assert request.schema is not None
        name = request.schema.__name__
        if name == "ClassifyResult":
            calls.append("classifier")
            tool_input = classify_payload
        elif name == "DocumentExtraction":
            calls.append("extractor_pass_a")
            assert pass_a_payload is not None
            tool_input = pass_a_payload
        else:  # pragma: no cover - defensive
            raise AssertionError(f"unexpected schema for anthropic transport: {name}")
        return TransportResponse(text="", tool_input=tool_input, input_tokens=10, output_tokens=10)

    def openai_transport(request: TransportRequest) -> TransportResponse:
        assert request.schema is not None
        calls.append("extractor_pass_b")
        assert pass_b_payload is not None
        return TransportResponse(
            text="", tool_input=pass_b_payload, input_tokens=10, output_tokens=10
        )

    bindings: dict[str, list[ModelBinding]] = {
        "classifier": [ModelBinding(provider="anthropic", model="fake-haiku")],
        "extractor_pass_a": [ModelBinding(provider="anthropic", model="fake-sonnet")],
        "extractor_pass_b": [ModelBinding(provider="openai", model="fake-gpt")],
    }
    providers = {
        "anthropic": AnthropicProvider(api_key=None, transport=anthropic_transport),
        "openai": OpenAIProvider(api_key=None, transport=openai_transport),
    }
    return LlmClient(bindings, providers)


def test_docx_narrative_document_becomes_a_full_text_encounter(tmp_path: Path) -> None:
    repo = DataRepo.init_at(tmp_path / "data-repo")
    db = LabsDb(tmp_path / "labs.sqlite")
    doc_path = tmp_path / "history.docx"
    _make_narrative_docx(doc_path)

    calls: list[str] = []
    client = _build_docx_llm_client(
        classify_payload={"doc_type": "clinical_note", "doc_date": "2026-08-01"},
        pass_a_payload=None,
        pass_b_payload=None,
        calls=calls,
    )
    vision = VisionClient(client)

    report = ingest_file(doc_path, repo=repo, db=db, vision=vision, clock=_fixed_clock)

    outcome = report.files[0]
    assert outcome.outcome == "ingested"
    assert outcome.doc_type == "clinical_note"
    assert outcome.rows_auto == 0
    assert outcome.rows_pending == 0
    assert outcome.commit_sha is not None
    # extractor roles are never called for a non-lab docx
    assert calls == ["classifier"]
    assert db.pending() == []

    documents = db.list_documents()
    assert len(documents) == 1
    assert documents[0].page_count == 1

    encounters_dir = repo.root / "case" / "encounters"
    encounter_files = list(encounters_dir.glob("2026-08-01--*.md"))
    assert len(encounter_files) == 1
    encounter = read_encounter(encounter_files[0])
    assert encounter.frontmatter.type == "patient-report"
    assert encounter.frontmatter.sources
    assert "Patient-authored clinical history." in encounter.extracted_text
    assert "Onset of joint pain in March 2026" in encounter.extracted_text

    git_repo = Repo(repo.root)
    assert git_repo.head.commit.hexsha == outcome.commit_sha


def test_docx_lab_report_double_pass_text_reconciles_to_auto_row(tmp_path: Path) -> None:
    repo = DataRepo.init_at(tmp_path / "data-repo")
    db = LabsDb(tmp_path / "labs.sqlite")
    doc_path = tmp_path / "home-lab.docx"
    _make_lab_docx(doc_path)

    result_row = {
        "name_raw": "Potassium",
        "value": 4.1,
        "unit_raw": "mmol/L",
        "ref_range_raw": "3.5-5.1",
        "page": 1,
        "confidence": "high",
    }
    extraction_payload = {
        "doc_type": "lab_report",
        "collection_date": "2026-08-01",
        "results": [result_row],
        "narrative_findings": [],
        "illegible_regions": [],
    }
    calls: list[str] = []
    client = _build_docx_llm_client(
        classify_payload={"doc_type": "lab_report", "doc_date": "2026-08-01"},
        pass_a_payload=extraction_payload,
        pass_b_payload=extraction_payload,
        calls=calls,
    )
    vision = VisionClient(client)

    report = ingest_file(doc_path, repo=repo, db=db, vision=vision, clock=_fixed_clock)

    outcome = report.files[0]
    assert outcome.outcome == "ingested"
    assert outcome.doc_type == "lab_report"
    assert outcome.rows_auto == 1
    assert outcome.rows_pending == 0
    assert outcome.issues == []
    assert calls == ["classifier", "extractor_pass_a", "extractor_pass_b"]

    documents = db.list_documents()
    assert len(documents) == 1
    assert documents[0].page_count == 1

    export_path = repo.root / "labs-export.jsonl"
    assert export_path.exists()
    lines = export_path.read_text(encoding="utf-8").strip().splitlines()
    assert any(json.loads(line)["table"] == "lab" for line in lines)
