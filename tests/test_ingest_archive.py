"""Tests for adoc.ingest.archive: sha256 archival, page rendering, dedupe."""

from __future__ import annotations

import shutil
import zipfile
from datetime import date, datetime
from pathlib import Path

import pytest
from conftest import TINY_PDF_BYTES, fake_page_renderer
from docx import Document

from adoc.ingest.archive import ArchiveError, archive_document, sha256_file
from adoc.labs.db import LabsDb
from adoc.labs.models import DocumentStatus, LabDocument


def test_sha256_file_is_deterministic(tiny_pdf_path: Path) -> None:
    assert sha256_file(tiny_pdf_path) == sha256_file(tiny_pdf_path)
    assert len(sha256_file(tiny_pdf_path)) == 64


def test_archive_document_copies_immutable_original_and_renders_pages(
    tmp_path: Path, tiny_pdf_path: Path
) -> None:
    repo_root = tmp_path / "data-repo"
    db = LabsDb(tmp_path / "labs.sqlite")

    archived = archive_document(repo_root, tiny_pdf_path, db=db, renderer=fake_page_renderer(3))

    assert archived.original_path.exists()
    assert archived.original_path.read_bytes() == TINY_PDF_BYTES
    assert archived.original_path.name == f"{archived.sha256}__document.pdf"
    assert archived.original_path.parent == repo_root / "sources"
    assert len(archived.page_paths) == 3
    for page_path in archived.page_paths:
        assert page_path.exists()
        assert page_path.parent == repo_root / "sources" / "pages" / archived.sha256
    assert archived.already_ingested is False


def test_archive_document_is_idempotent_and_does_not_re_render(
    tmp_path: Path, tiny_pdf_path: Path
) -> None:
    repo_root = tmp_path / "data-repo"
    db = LabsDb(tmp_path / "labs.sqlite")
    calls = {"count": 0}

    def counting_renderer(pdf_path: Path, out_dir: Path) -> list[Path]:
        calls["count"] += 1
        return fake_page_renderer(2)(pdf_path, out_dir)

    first = archive_document(repo_root, tiny_pdf_path, db=db, renderer=counting_renderer)
    second = archive_document(repo_root, tiny_pdf_path, db=db, renderer=counting_renderer)

    assert calls["count"] == 1
    assert first.sha256 == second.sha256
    assert first.page_paths == second.page_paths


def test_archive_document_reports_already_ingested_from_labs_db(
    tmp_path: Path, tiny_pdf_path: Path
) -> None:
    repo_root = tmp_path / "data-repo"
    db = LabsDb(tmp_path / "labs.sqlite")
    sha = sha256_file(tiny_pdf_path)
    db.upsert_document(
        LabDocument(
            sha256=sha,
            filename="document.pdf",
            doc_type="lab_report",
            doc_date=date(2026, 5, 2),
            page_count=1,
            ingested_at=datetime(2026, 5, 3, 0, 0, 0),
            status=DocumentStatus.COMPLETE,
        )
    )

    archived = archive_document(repo_root, tiny_pdf_path, db=db, renderer=fake_page_renderer(1))

    assert archived.already_ingested is True


def test_pdftoppm_renderer_raises_clear_error_when_binary_missing(
    tmp_path: Path, tiny_pdf_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(shutil, "which", lambda _name: None)

    db = LabsDb(tmp_path / "labs.sqlite")
    with pytest.raises(ArchiveError, match="pdftoppm"):
        archive_document(tmp_path / "data-repo", tiny_pdf_path, db=db)


def test_archive_rejects_non_pdf_files(tmp_path: Path) -> None:
    """Junk/fake-docx must never reach the immutable sources/ store."""
    import pytest

    from adoc.ingest.archive import ArchiveError, archive_document

    bogus = tmp_path / "notes.docx"
    bogus.write_bytes(b"PK\x03\x04 not a pdf")

    with pytest.raises(ArchiveError, match="not a PDF"):
        archive_document(tmp_path, bogus, db=None, renderer=None)  # type: ignore[arg-type]


def _make_docx(path: Path, text: str = "hello") -> None:
    document = Document()
    document.add_paragraph(text)
    document.save(str(path))


def test_archive_document_archives_a_real_docx_with_no_page_rendering(tmp_path: Path) -> None:
    repo_root = tmp_path / "data-repo"
    db = LabsDb(tmp_path / "labs.sqlite")
    docx_path = tmp_path / "history.docx"
    _make_docx(docx_path, "Patient-authored clinical history.")

    def exploding_renderer(pdf_path: Path, out_dir: Path) -> list[Path]:  # pragma: no cover
        raise AssertionError("a docx must never be rendered to page images")

    archived = archive_document(repo_root, docx_path, db=db, renderer=exploding_renderer)

    assert archived.kind == "docx"
    assert archived.page_paths == []
    assert archived.original_path.exists()
    assert archived.original_path.read_bytes() == docx_path.read_bytes()
    assert archived.original_path.name == f"{archived.sha256}__history.docx"
    assert archived.already_ingested is False


def test_archive_document_dedupes_docx_like_pdf(tmp_path: Path) -> None:
    repo_root = tmp_path / "data-repo"
    db = LabsDb(tmp_path / "labs.sqlite")
    docx_path = tmp_path / "history.docx"
    _make_docx(docx_path)

    first = archive_document(repo_root, docx_path, db=db, renderer=fake_page_renderer(0))
    db.upsert_document(
        LabDocument(
            sha256=first.sha256,
            filename="history.docx",
            doc_type="clinical_note",
            page_count=1,
            status=DocumentStatus.NEEDS_REVIEW,
        )
    )
    second = archive_document(repo_root, docx_path, db=db, renderer=fake_page_renderer(0))

    assert first.sha256 == second.sha256
    assert second.already_ingested is True
    assert second.kind == "docx"
    assert second.page_paths == []


# --------------------------------------------------------------------------
# text (genomics/filetypes task item 1/3): archived like docx - no page
# rendering, size-capped rather than silently truncated.
# --------------------------------------------------------------------------


def test_archive_document_archives_a_plain_text_document_with_no_page_rendering(
    tmp_path: Path,
) -> None:
    repo_root = tmp_path / "data-repo"
    db = LabsDb(tmp_path / "labs.sqlite")
    text_path = tmp_path / "history.txt"
    text_path.write_text("Patient-authored clinical history.", encoding="utf-8")

    def exploding_renderer(pdf_path: Path, out_dir: Path) -> list[Path]:  # pragma: no cover
        raise AssertionError("a text document must never be rendered to page images")

    archived = archive_document(repo_root, text_path, db=db, renderer=exploding_renderer)

    assert archived.kind == "text"
    assert archived.page_paths == []
    assert archived.original_path.read_bytes() == text_path.read_bytes()
    assert archived.original_path.name == f"{archived.sha256}__history.txt"


def test_archive_document_rejects_oversized_text_with_a_clear_reason(tmp_path: Path) -> None:
    """Larger non-genomic text is rejected with a clear reason rather than
    silently truncated - and never copied into `sources/` at all."""
    repo_root = tmp_path / "data-repo"
    db = LabsDb(tmp_path / "labs.sqlite")
    text_path = tmp_path / "huge.txt"
    text_path.write_text("x" * 100, encoding="utf-8")

    with pytest.raises(ArchiveError, match="larger than the"):
        archive_document(repo_root, text_path, db=db, text_max_bytes=50)

    assert not any((repo_root / "sources").glob("*__huge.txt"))


def test_archive_document_rejects_genomic_and_zip_kinds(tmp_path: Path) -> None:
    """`archive_document` is never the path for genomic files or zip
    archives - `ingest.pipeline` routes those elsewhere before archival is
    ever attempted; calling it directly on either is a programming error."""
    repo_root = tmp_path / "data-repo"
    db = LabsDb(tmp_path / "labs.sqlite")

    vcf_path = tmp_path / "sample.vcf"
    vcf_path.write_bytes(b"##fileformat=VCFv4.2\n#CHROM\tPOS\n1\t1\n")
    with pytest.raises(ArchiveError, match="dedicated path"):
        archive_document(repo_root, vcf_path, db=db)

    zip_path = tmp_path / "bundle.zip"
    with zipfile.ZipFile(zip_path, "w") as archive:
        archive.writestr("a.txt", "hello")
    with pytest.raises(ArchiveError, match="dedicated path"):
        archive_document(repo_root, zip_path, db=db)
